"""
DOCX Generator for document export.

Produces Word documents from structured content (title, sections with headings,
paragraphs, bullet lists, tables). Used by the Documents API export endpoint.

Reference: docs/stages/STAGE_WORKER_SPECS.md Stage 16/19/20.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any, Dict


def generate_docx(content: dict) -> bytes:
    """
    Generate a DOCX file from structured content.

    Accepts:
        content: {
            "title": str,
            "sections": [
                {
                    "heading": str,
                    "paragraphs": [str],   # optional
                    "bullets": [str],      # optional
                    "table": [[str]]       # optional, list of rows
                }
            ]
        }

    Returns:
        DOCX file as bytes.
    """
    from docx import Document

    doc = Document()
    title = (content.get("title") or "").strip() or "Untitled"
    doc.add_heading(title, level=0)

    sections = content.get("sections") or []
    for sec in sections:
        if not isinstance(sec, dict):
            continue
        heading = (sec.get("heading") or "").strip() or "Section"
        doc.add_heading(heading, level=1)

        # Paragraphs
        paragraphs = sec.get("paragraphs")
        if paragraphs:
            for p in paragraphs:
                if p is None:
                    continue
                text = str(p).strip()
                doc.add_paragraph(text or "")

        # Bullet list
        bullets = sec.get("bullets")
        if bullets:
            for b in bullets:
                if b is None:
                    continue
                text = str(b).strip()
                doc.add_paragraph(text, style="List Bullet")

        # Table
        table_data = sec.get("table")
        if table_data and isinstance(table_data, list) and len(table_data) > 0:
            rows = [r for r in table_data if isinstance(r, list)]
            if rows:
                col_count = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=col_count)
                table.style = "Table Grid"
                for i, row in enumerate(rows):
                    for j, cell in enumerate(row):
                        if j < col_count:
                            table.rows[i].cells[j].text = str(cell) if cell is not None else ""

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
